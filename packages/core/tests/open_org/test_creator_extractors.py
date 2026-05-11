"""Tests for document text extraction used by the chat creator's upload step.

We support PDF, DOCX, plain text. Other file types raise. A hard size cap
protects the worker process from a malicious or accidental large upload.
"""

from __future__ import annotations

import io
from unittest import mock

import pytest

from llmstxt_core.open_org.creator.extractors import (
    MAX_UPLOAD_BYTES,
    ExtractError,
    UnsupportedFormatError,
    extract_text,
)


# --- Plain text ------------------------------------------------------------


def test_extract_plain_text():
    text = extract_text("notes.txt", b"Hello from Open Org Creator")
    assert text == "Hello from Open Org Creator"


def test_extract_plain_text_decodes_utf8():
    text = extract_text("notes.txt", "Strategy fünding for café".encode("utf-8"))
    assert "fünding" in text
    assert "café" in text


def test_extract_plain_text_falls_back_to_latin1_on_bad_utf8():
    """Older docs sometimes arrive as Latin-1. We don't want to error out."""
    text = extract_text("notes.txt", b"\xa3 100k for the trust")
    # The pound sign survives one way or the other.
    assert "100k for the trust" in text


# --- DOCX (real round-trip) ------------------------------------------------


def _make_docx_bytes(text: str) -> bytes:
    """Build a real .docx in memory using python-docx."""
    from docx import Document  # type: ignore[import-not-found]

    document = Document()
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_docx_returns_paragraph_text():
    docx_bytes = _make_docx_bytes("Riverside Community Trust supports older people.")
    text = extract_text("strategy.docx", docx_bytes)
    assert "Riverside Community Trust" in text
    assert "older people" in text


def test_extract_docx_joins_paragraphs_with_newlines():
    from docx import Document

    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text("strategy.docx", buffer.getvalue())
    assert text.count("\n") >= 1
    assert "First paragraph." in text
    assert "Second paragraph." in text


# --- PDF (mocked) ----------------------------------------------------------


def _fake_pdf_reader(pages_text: list[str]):
    class _Page:
        def __init__(self, t):
            self._t = t

        def extract_text(self):
            return self._t

    reader = mock.MagicMock()
    reader.pages = [_Page(t) for t in pages_text]
    return reader


def test_extract_pdf_concatenates_pages():
    with mock.patch(
        "llmstxt_core.open_org.creator.extractors.PdfReader",
        return_value=_fake_pdf_reader(
            ["Page one content.", "Page two content.", ""]
        ),
    ):
        text = extract_text("strategy.pdf", b"\x25PDF-fake")

    assert "Page one content." in text
    assert "Page two content." in text


def test_extract_pdf_handles_pypdf_error_gracefully():
    """Bad PDFs raise ExtractError rather than leaking the underlying exception."""
    with mock.patch(
        "llmstxt_core.open_org.creator.extractors.PdfReader",
        side_effect=Exception("PDF stream is corrupted"),
    ):
        with pytest.raises(ExtractError):
            extract_text("strategy.pdf", b"not really a pdf")


# --- Size cap + format gate ------------------------------------------------


def test_rejects_oversize_upload():
    too_big = b"x" * (MAX_UPLOAD_BYTES + 1)
    with pytest.raises(ExtractError, match="too large"):
        extract_text("notes.txt", too_big)


def test_rejects_unsupported_format():
    with pytest.raises(UnsupportedFormatError):
        extract_text("strategy.rtf", b"some content")


def test_rejects_empty_content():
    with pytest.raises(ExtractError, match="empty"):
        extract_text("notes.txt", b"")


def test_dispatch_is_case_insensitive_on_extension():
    text = extract_text("STRATEGY.TXT", b"hello")
    assert text == "hello"
